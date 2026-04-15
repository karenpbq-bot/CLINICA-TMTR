"""
Generador de archivos Word (.docx) y Excel (.xlsx) para reportes de ORTHOCLINIC.
Librerías: python-docx, pandas, openpyxl.
"""

import os
import datetime
from pathlib import Path

_LOGO_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                           "assets", "logo_orthoclinic.jpg")
_OUT_DIR   = os.path.join(os.path.dirname(os.path.abspath(__file__)), "reportes")

os.makedirs(_OUT_DIR, exist_ok=True)

ANTECEDENTES = [
    ("tratamiento_medico",       "01. Tratamiento médico actual"),
    ("medicamentos",             "02. Toma de medicamentos"),
    ("alergias_med",             "03. Alergias a medicamentos"),
    ("cardiopatias",             "04. Cardiopatías"),
    ("presion_arterial_alt",     "05. Alteración presión arterial"),
    ("embarazo",                 "06. Embarazo"),
    ("diabetes",                 "07. Diabetes"),
    ("hepatitis",                "08. Hepatitis"),
    ("irradiaciones",            "09. Irradiaciones previas"),
    ("discrasias",               "10. Discrasias sanguíneas"),
    ("fiebre_reumatica",         "11. Fiebre reumática"),
    ("enf_renales",              "12. Enfermedades renales"),
    ("inmunosupresion",          "13. Inmunosupresión / VIH"),
    ("trastornos_emocionales",   "14. Trastornos emocionales"),
    ("trastornos_respiratorios", "15. Trastornos respiratorios"),
    ("trastornos_gastricos",     "16. Trastornos gástricos"),
    ("epilepsia",                "17. Epilepsia"),
    ("cirugias",                 "18. Cirugías previas"),
    ("enf_orales",               "19. Enfermedades orales previas"),
    ("otras_alteraciones",       "20. Otras alteraciones sistémicas"),
    ("fuma_licor",               "21. Tabaquismo / Consumo de alcohol"),
]


# ═══════════════════════════════════════════════════════════════════════════════
#  Helpers comunes
# ═══════════════════════════════════════════════════════════════════════════════

def _fmt(v, default="—") -> str:
    return str(v).strip() if v else default


def _fmt_fecha(iso: str) -> str:
    if not iso:
        return "—"
    try:
        return datetime.datetime.fromisoformat(
            iso.replace("Z", "+00:00")).strftime("%d/%m/%Y")
    except Exception:
        return str(iso)[:10]


def _fmt_monto(v) -> str:
    try:
        return f"$ {float(v):,.2f}"
    except Exception:
        return "$ 0,00"


def _safe_name(texto: str) -> str:
    return "".join(c if c.isalnum() or c in "._- " else "_" for c in texto)[:60]


# ═══════════════════════════════════════════════════════════════════════════════
#  WORD — Historia Clínica
# ═══════════════════════════════════════════════════════════════════════════════

def generar_historia_clinica_docx(paciente: dict, historia: dict,
                                   output_dir: str | None = None) -> str:
    """
    Genera un documento Word con ficha del paciente, anamnesis y constantes vitales.
    Devuelve la ruta al archivo generado.
    """
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL

    output_dir = output_dir or _OUT_DIR
    os.makedirs(output_dir, exist_ok=True)

    apellido = _fmt(paciente.get("apellido", "")).upper()
    nombre   = _fmt(paciente.get("nombre", ""))
    safe     = _safe_name(f"HC_{apellido}_{nombre}")
    ruta     = os.path.join(output_dir, f"{safe}.docx")

    doc = Document()

    # ── Márgenes ──────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin   = Cm(2.2)
        section.right_margin  = Cm(2.2)

    AZUL = RGBColor(0x15, 0x65, 0xC0)
    GRIS = RGBColor(0x61, 0x61, 0x61)

    # ── Encabezado con logo ───────────────────────────────────────────────
    hdr_table = doc.add_table(rows=1, cols=2)
    hdr_table.style = "Table Grid"
    cell_logo = hdr_table.cell(0, 0)
    cell_info = hdr_table.cell(0, 1)
    cell_logo.width = Cm(4)

    if os.path.exists(_LOGO_PATH):
        from docx.shared import Inches
        p = cell_logo.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run()
        r.add_picture(_LOGO_PATH, width=Cm(3.5))

    p_info = cell_info.paragraphs[0]
    p_info.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p_info.add_run("ORTHOCLINIC\n")
    run.bold = True; run.font.size = Pt(16); run.font.color.rgb = AZUL
    run2 = p_info.add_run("Historia Clínica Odontológica\n")
    run2.font.size = Pt(11); run2.font.color.rgb = GRIS
    run3 = p_info.add_run(f"Generado: {datetime.datetime.today().strftime('%d/%m/%Y %H:%M')}")
    run3.font.size = Pt(9); run3.font.color.rgb = GRIS

    # Eliminar bordes del encabezado
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    def _sin_bordes(tabla):
        for row in tabla.rows:
            for cell in row.cells:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcBorders = OxmlElement("w:tcBorders")
                for side in ("top","left","bottom","right","insideH","insideV"):
                    bd = OxmlElement(f"w:{side}")
                    bd.set(qn("w:val"), "none")
                    tcBorders.append(bd)
                tcPr.append(tcBorders)
    _sin_bordes(hdr_table)

    doc.add_paragraph()

    # ── Título paciente ───────────────────────────────────────────────────
    def _seccion(titulo: str):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(4)
        run = p.add_run(f"  {titulo}  ")
        run.bold = True; run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.runs[0].font.highlight_color = None
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "1565C0")
        p._p.get_or_add_pPr().append(shd)
        return p

    def _campo_tabla(tabla, fila, col, lbl, val):
        c = tabla.cell(fila, col)
        p = c.paragraphs[0]
        r1 = p.add_run(lbl + ": ")
        r1.bold = True; r1.font.size = Pt(9.5)
        r2 = p.add_run(_fmt(val))
        r2.font.size = Pt(9.5)

    # ── SECCIÓN 1: FICHA DEL PACIENTE ─────────────────────────────────────
    _seccion("1. FICHA DEL PACIENTE")

    tbl = doc.add_table(rows=4, cols=4)
    tbl.style = "Table Grid"
    for row in tbl.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    _campo_tabla(tbl, 0, 0, "Apellido",      paciente.get("apellido", ""))
    _campo_tabla(tbl, 0, 1, "Nombre",        paciente.get("nombre", ""))
    _campo_tabla(tbl, 0, 2, "DNI / Cédula",  paciente.get("dni", ""))
    _campo_tabla(tbl, 0, 3, "Fec. nacim.",   _fmt_fecha(paciente.get("fecha_nac", "")))
    _campo_tabla(tbl, 1, 0, "Grupo sang.",   paciente.get("grupo_sangre", ""))
    _campo_tabla(tbl, 1, 1, "Teléfono",      paciente.get("telefono", ""))
    _campo_tabla(tbl, 1, 2, "Email",         paciente.get("email", ""))
    _campo_tabla(tbl, 1, 3, "Obra Social",   paciente.get("obra_social", ""))
    _campo_tabla(tbl, 2, 0, "Nro. afiliado", paciente.get("nro_afiliado", ""))
    # Mergar celdas dirección
    tbl.cell(2, 1).merge(tbl.cell(2, 2)).merge(tbl.cell(2, 3))
    _campo_tabla(tbl, 2, 1, "Dirección",     paciente.get("direccion", ""))
    tbl.cell(3, 0).merge(tbl.cell(3, 1)).merge(tbl.cell(3, 2)).merge(tbl.cell(3, 3))
    _campo_tabla(tbl, 3, 0, "Alergias conocidas", paciente.get("alergias", ""))

    doc.add_paragraph()

    # ── SECCIÓN 2: ANAMNESIS ──────────────────────────────────────────────
    _seccion("2. ANAMNESIS — Antecedentes médicos y odontológicos")
    ant_dict = (historia.get("antecedentes") or {})

    tbl_ant = doc.add_table(rows=0, cols=3)
    tbl_ant.style = "Table Grid"
    ITEMS_POR_FILA = 3
    fila_actual = None
    for i, (key, label) in enumerate(ANTECEDENTES):
        if i % ITEMS_POR_FILA == 0:
            fila_actual = tbl_ant.add_row()
        positivo = bool(ant_dict.get(key, False))
        simbolo  = "✅ SÍ" if positivo else "○  NO"
        col      = i % ITEMS_POR_FILA
        p        = fila_actual.cells[col].paragraphs[0]
        r1       = p.add_run(f"{simbolo}  ")
        r1.bold  = positivo; r1.font.size = Pt(9)
        if positivo:
            r1.font.color.rgb = RGBColor(0xC6, 0x28, 0x28)
        r2       = p.add_run(label)
        r2.font.size = Pt(9)

    doc.add_paragraph()

    # ── SECCIÓN 3: CONSTANTES VITALES ────────────────────────────────────
    _seccion("3. CONSTANTES VITALES")
    sv = historia.get("signos_vitales") or {}

    try:
        p_peso = float((sv.get("peso","0") or "0").replace(",","."))
        p_est  = float((sv.get("estatura","0") or "0").replace(",","."))
        imc    = f"{p_peso/(p_est/100)**2:.1f}" if p_peso > 0 and p_est > 0 else "—"
    except Exception:
        imc = "—"

    tbl_sv = doc.add_table(rows=2, cols=7)
    tbl_sv.style = "Table Grid"
    cabeceras = ["Presión", "Pulso", "Temperatura", "F. Resp.", "Peso", "Talla", "IMC"]
    valores   = [
        sv.get("tension_arterial",""),
        sv.get("pulso",""),
        sv.get("temperatura",""),
        sv.get("frecuencia_resp",""),
        (sv.get("peso","") + " kg") if sv.get("peso") else "—",
        (sv.get("estatura","") + " cm") if sv.get("estatura") else "—",
        imc,
    ]
    for i, (cab, val) in enumerate(zip(cabeceras, valores)):
        c = tbl_sv.cell(0, i).paragraphs[0]
        r = c.add_run(cab)
        r.bold = True; r.font.size = Pt(9)
        r.font.color.rgb = AZUL
        c2 = tbl_sv.cell(1, i).paragraphs[0]
        r2 = c2.add_run(_fmt(val))
        r2.font.size = Pt(9.5)

    doc.add_paragraph()

    # ── SECCIÓN 4: DATOS DE LA CONSULTA ──────────────────────────────────
    _seccion("4. DATOS DE LA CONSULTA")

    tbl_c = doc.add_table(rows=2, cols=3)
    tbl_c.style = "Table Grid"
    _campo_tabla(tbl_c, 0, 0, "N° Historia",   historia.get("historia_no",""))
    _campo_tabla(tbl_c, 0, 1, "Odontólogo",    historia.get("odontologo",""))
    _campo_tabla(tbl_c, 0, 2, "Fecha",         _fmt_fecha(historia.get("fecha_elaboracion","")))

    tbl_c.cell(1, 0).merge(tbl_c.cell(1, 1)).merge(tbl_c.cell(1, 2))

    for lbl, key in [
        ("Motivo de consulta",                    "motivo_consulta"),
        ("Enfermedad actual / Hallazgos clínicos","enfermedad_actual"),
        ("Observaciones y plan de tratamiento",   "observaciones"),
    ]:
        doc.add_paragraph()
        p_lbl = doc.add_paragraph()
        r_lbl = p_lbl.add_run(lbl.upper() + ":")
        r_lbl.bold = True; r_lbl.font.size = Pt(9.5); r_lbl.font.color.rgb = AZUL
        val   = (historia.get(key) or "").strip() or "—"
        p_val = doc.add_paragraph(val)
        p_val.paragraph_format.left_indent = Cm(0.5)
        for run in p_val.runs:
            run.font.size = Pt(9.5)

    # ── Pie ───────────────────────────────────────────────────────────────
    doc.add_paragraph()
    pie = doc.add_paragraph()
    pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_pie = pie.add_run(
        f"Documento confidencial — uso exclusivo del profesional · "
        f"Generado el {datetime.datetime.today().strftime('%d/%m/%Y %H:%M')}"
    )
    r_pie.font.size = Pt(8); r_pie.font.color.rgb = GRIS

    doc.save(ruta)
    return ruta


# ═══════════════════════════════════════════════════════════════════════════════
#  EXCEL — Presupuestos / Resumen Financiero
# ═══════════════════════════════════════════════════════════════════════════════

def generar_excel_presupuestos(datos: list[dict],
                                filtros_desc: str = "",
                                output_dir: str | None = None) -> str:
    """
    Genera Excel con resumen financiero de tratamientos.
    datos: lista de tratamientos con 'pagado' y 'saldo' calculados.
    """
    import openpyxl
    from openpyxl.styles import (Font, PatternFill, Alignment,
                                  Border, Side, numbers)
    from openpyxl.utils import get_column_letter

    output_dir = output_dir or _OUT_DIR
    os.makedirs(output_dir, exist_ok=True)
    hoy   = datetime.datetime.today()
    ruta  = os.path.join(output_dir,
                         f"Presupuestos_{hoy.strftime('%Y%m%d_%H%M')}.xlsx")

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Resumen Financiero"

    AZUL_F = "1565C0"; AZUL_BG = "E3F2FD"; GRIS_BG = "FAFAFA"
    VERDE   = "1B5E20"; VERDE_BG = "E8F5E9"
    ROJO_BG = "FFEBEE"; ROJO_F  = "C62828"

    thin = Side(style="thin", color="BDBDBD")
    brd  = Border(left=thin, right=thin, top=thin, bottom=thin)

    def _celda(ws, fila, col, valor, bold=False, bg=None, fg="212121",
               align="left", num_fmt=None):
        c = ws.cell(row=fila, column=col, value=valor)
        c.font      = Font(bold=bold, color=fg, name="Calibri", size=10)
        c.alignment = Alignment(horizontal=align, vertical="center", wrap_text=True)
        c.border    = brd
        if bg:
            c.fill  = PatternFill("solid", fgColor=bg)
        if num_fmt:
            c.number_format = num_fmt
        return c

    # ── Logo + título ──────────────────────────────────────────────────────
    ws.merge_cells("A1:J2")
    c_titulo = ws["A1"]
    c_titulo.value     = "ORTHOCLINIC — Resumen Financiero de Presupuestos"
    c_titulo.font      = Font(bold=True, size=14, color="FFFFFF", name="Calibri")
    c_titulo.fill      = PatternFill("solid", fgColor=AZUL_F)
    c_titulo.alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[1].height = 24
    ws.row_dimensions[2].height = 24

    ws.merge_cells("A3:J3")
    c_sub = ws["A3"]
    c_sub.value     = (
        f"Generado: {hoy.strftime('%d/%m/%Y %H:%M')}"
        + (f"   |   Filtros: {filtros_desc}" if filtros_desc else "")
    )
    c_sub.font      = Font(italic=True, size=9, color="616161", name="Calibri")
    c_sub.alignment = Alignment(horizontal="left", vertical="center")

    # ── Encabezado tabla ──────────────────────────────────────────────────
    COLS = [
        ("Paciente",     22), ("DNI",          13),
        ("Obra Social",  16), ("Especialista",  20),
        ("Descripción",  30), ("Diente",         8),
        ("Estado",       14), ("Costo",          12),
        ("Pagado",       12), ("Saldo",          12),
    ]
    HDR_ROW = 4
    for i, (lbl, ancho) in enumerate(COLS, 1):
        _celda(ws, HDR_ROW, i, lbl, bold=True, bg=AZUL_F, fg="FFFFFF", align="center")
        ws.column_dimensions[get_column_letter(i)].width = ancho
    ws.row_dimensions[HDR_ROW].height = 20

    # ── Datos ─────────────────────────────────────────────────────────────
    _ESTADO_BG = {
        "presupuestado": "FFF8E1",
        "aprobado":      "E8F5E9",
        "realizado":     "E3F2FD",
    }
    total_costo = total_pagado = total_saldo = 0.0
    for i, t in enumerate(datos):
        fila = HDR_ROW + 1 + i
        bg   = _ESTADO_BG.get(t.get("estado",""), "FFFFFF")
        pac  = t.get("pacientes") or {}
        esp  = t.get("especialistas") or {}
        nom_pac = f"{_fmt(pac.get('apellido',''))} {_fmt(pac.get('nombre',''))}".strip()
        nom_esp = f"{_fmt(esp.get('apellido',''))} {_fmt(esp.get('nombre',''))}".strip()
        costo  = float(t.get("costo", 0))
        pagado = float(t.get("pagado", 0))
        saldo  = float(t.get("saldo", 0))
        total_costo  += costo
        total_pagado += pagado
        total_saldo  += saldo

        _celda(ws, fila, 1, nom_pac,                 bg=bg)
        _celda(ws, fila, 2, _fmt(pac.get("dni","")), bg=bg)
        _celda(ws, fila, 3, _fmt(pac.get("obra_social","")), bg=bg)
        _celda(ws, fila, 4, nom_esp,                 bg=bg)
        _celda(ws, fila, 5, _fmt(t.get("descripcion","")), bg=bg)
        _celda(ws, fila, 6, str(t.get("diente","")) if t.get("diente") else "—",
               bg=bg, align="center")
        _celda(ws, fila, 7, (t.get("estado","")).capitalize(), bg=bg, align="center")
        _celda(ws, fila, 8, costo,  bg=bg, align="right", num_fmt='"$"#,##0.00')
        _celda(ws, fila, 9, pagado, bg=bg, align="right", num_fmt='"$"#,##0.00')
        sf = saldo
        _celda(ws, fila, 10, saldo,
               bg=ROJO_BG if saldo > 0 else VERDE_BG,
               fg=ROJO_F  if saldo > 0 else VERDE,
               align="right", num_fmt='"$"#,##0.00')
        ws.row_dimensions[fila].height = 16

    # ── Fila totales ──────────────────────────────────────────────────────
    f_tot = HDR_ROW + 1 + len(datos)
    ws.merge_cells(f"A{f_tot}:G{f_tot}")
    _celda(ws, f_tot, 1, f"TOTALES  ({len(datos)} tratamientos)",
           bold=True, bg=AZUL_F, fg="FFFFFF", align="right")
    _celda(ws, f_tot, 8, total_costo,  bold=True, bg=AZUL_BG,
           align="right", num_fmt='"$"#,##0.00')
    _celda(ws, f_tot, 9, total_pagado, bold=True, bg=VERDE_BG,
           fg=VERDE, align="right", num_fmt='"$"#,##0.00')
    _celda(ws, f_tot, 10, total_saldo, bold=True,
           bg=ROJO_BG if total_saldo > 0 else VERDE_BG,
           fg=ROJO_F  if total_saldo > 0 else VERDE,
           align="right", num_fmt='"$"#,##0.00')
    ws.row_dimensions[f_tot].height = 20

    # Fijar encabezado
    ws.freeze_panes = f"A{HDR_ROW + 1}"

    wb.save(ruta)
    return ruta


# ═══════════════════════════════════════════════════════════════════════════════
#  EXCEL — Agenda / Cronograma de Citas
# ═══════════════════════════════════════════════════════════════════════════════

def generar_excel_agenda(datos: list[dict], especialista_nombre: str = "",
                          periodo: str = "semana",
                          output_dir: str | None = None) -> str:
    """
    Genera Excel con cronograma de citas del especialista.
    """
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter

    output_dir = output_dir or _OUT_DIR
    os.makedirs(output_dir, exist_ok=True)
    hoy  = datetime.datetime.today()
    safe = _safe_name(especialista_nombre or "Todos")
    ruta = os.path.join(
        output_dir,
        f"Agenda_{safe}_{periodo}_{hoy.strftime('%Y%m%d')}.xlsx"
    )

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Cronograma"

    AZUL_F  = "1565C0"
    thin    = Side(style="thin", color="BDBDBD")
    brd     = Border(left=thin, right=thin, top=thin, bottom=thin)
    DIAS_ES = {0:"Lunes",1:"Martes",2:"Miércoles",3:"Jueves",
               4:"Viernes",5:"Sábado",6:"Domingo"}

    def _celda(ws, fila, col, valor, bold=False, bg=None, fg="212121",
               align="left", size=10):
        c = ws.cell(row=fila, column=col, value=valor)
        c.font      = Font(bold=bold, color=fg, name="Calibri", size=size)
        c.alignment = Alignment(horizontal=align, vertical="center", wrap_text=True)
        c.border    = brd
        if bg:
            c.fill  = PatternFill("solid", fgColor=bg)
        return c

    # ── Título ────────────────────────────────────────────────────────────
    PERIODOS = {"semana":"Próxima semana (7 días)",
                "quincena":"Próxima quincena (15 días)",
                "mes":"Próximo mes (30 días)"}
    periodo_txt = PERIODOS.get(periodo, periodo)
    ws.merge_cells("A1:I2")
    c_t = ws["A1"]
    c_t.value     = f"ORTHOCLINIC — Cronograma de Citas"
    c_t.font      = Font(bold=True, size=14, color="FFFFFF", name="Calibri")
    c_t.fill      = PatternFill("solid", fgColor=AZUL_F)
    c_t.alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[1].height = 22
    ws.row_dimensions[2].height = 22

    ws.merge_cells("A3:I3")
    c_s = ws["A3"]
    esp_txt = f"Especialista: {especialista_nombre}" if especialista_nombre else "Todos los especialistas"
    c_s.value = f"{esp_txt}   |   Período: {periodo_txt}   |   Generado: {hoy.strftime('%d/%m/%Y %H:%M')}"
    c_s.font  = Font(italic=True, size=9, color="616161", name="Calibri")
    c_s.alignment = Alignment(horizontal="left", vertical="center")

    # ── Encabezado ────────────────────────────────────────────────────────
    COLS = [
        ("Fecha",       12), ("Día",          10),
        ("Hora",         8), ("Paciente",     24),
        ("Teléfono",    14), ("Especialista", 22),
        ("Motivo",      28), ("Estado",       14),
        ("Notas",       22),
    ]
    HDR_ROW = 4
    for i, (lbl, ancho) in enumerate(COLS, 1):
        _celda(ws, HDR_ROW, i, lbl, bold=True, bg=AZUL_F, fg="FFFFFF", align="center")
        ws.column_dimensions[get_column_letter(i)].width = ancho
    ws.row_dimensions[HDR_ROW].height = 20

    # ── Datos ─────────────────────────────────────────────────────────────
    _EST_BG = {
        "pendiente":  "FFF8E1", "confirmada": "E8F5E9",
        "realizada":  "E3F2FD", "cancelada":  "FFEBEE",
    }
    _EST_FG = {
        "pendiente":  "E65100", "confirmada": "1B5E20",
        "realizada":  "1565C0", "cancelada":  "C62828",
    }

    for i, c in enumerate(datos):
        fila   = HDR_ROW + 1 + i
        pac    = c.get("pacientes") or {}
        esp    = c.get("especialistas") or {}
        estado = c.get("estado", "")
        bg     = _EST_BG.get(estado, "FFFFFF")
        fg     = _EST_FG.get(estado, "212121")
        nom_pac = f"{_fmt(pac.get('apellido',''))} {_fmt(pac.get('nombre',''))}".strip()
        nom_esp = f"{_fmt(esp.get('apellido',''))} {_fmt(esp.get('nombre',''))}".strip()
        iso     = c.get("fecha_hora","")
        try:
            dt    = datetime.datetime.fromisoformat(iso.replace("Z","+00:00"))
            fecha = dt.strftime("%d/%m/%Y")
            dia   = DIAS_ES.get(dt.weekday(), "")
            hora  = dt.strftime("%H:%M")
        except Exception:
            fecha = iso[:10]; dia = ""; hora = ""

        _celda(ws, fila, 1, fecha,   bg=bg, align="center")
        _celda(ws, fila, 2, dia,     bg=bg, align="center")
        _celda(ws, fila, 3, hora,    bg=bg, align="center")
        _celda(ws, fila, 4, nom_pac, bg=bg)
        _celda(ws, fila, 5, _fmt(pac.get("telefono","")), bg=bg)
        _celda(ws, fila, 6, nom_esp, bg=bg)
        _celda(ws, fila, 7, _fmt(c.get("motivo","")),    bg=bg)
        _celda(ws, fila, 8, estado.capitalize(), bold=True,
               bg=bg, fg=fg, align="center")
        _celda(ws, fila, 9, _fmt(c.get("notas","")), bg=bg)
        ws.row_dimensions[fila].height = 16

    # ── Total ─────────────────────────────────────────────────────────────
    f_tot = HDR_ROW + 1 + len(datos)
    ws.merge_cells(f"A{f_tot}:I{f_tot}")
    c_tot = ws.cell(row=f_tot, column=1,
                    value=f"Total: {len(datos)} cita(s) — {periodo_txt}")
    c_tot.font      = Font(bold=True, size=10, color="FFFFFF", name="Calibri")
    c_tot.fill      = PatternFill("solid", fgColor=AZUL_F)
    c_tot.alignment = Alignment(horizontal="right", vertical="center")
    c_tot.border    = brd
    ws.row_dimensions[f_tot].height = 18

    ws.freeze_panes = f"A{HDR_ROW + 1}"
    wb.save(ruta)
    return ruta


# ═══════════════════════════════════════════════════════════════════════════════
#  EXCEL — Reporte Financiero Total (multi-paciente / multi-especialista)
# ═══════════════════════════════════════════════════════════════════════════════

def generar_excel_financiero_total(datos: list[dict],
                                    filtros_desc: str = "",
                                    saldo_minimo: float = 0.0,
                                    output_dir: str | None = None) -> str:
    """
    Genera Excel con resumen financiero agrupado por paciente+especialista.
    datos: resultado de obtener_reporte_financiero_total().
    """
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter

    output_dir = output_dir or _OUT_DIR
    os.makedirs(output_dir, exist_ok=True)
    hoy  = datetime.datetime.today()
    ruta = os.path.join(output_dir,
                        f"Financiero_{hoy.strftime('%Y%m%d_%H%M')}.xlsx")

    wb  = openpyxl.Workbook()
    ws  = wb.active
    ws.title = "Reporte Financiero"

    AZUL_F  = "1565C0"; AZUL_BG = "E3F2FD"
    VERDE   = "1B5E20"; VERDE_BG = "E8F5E9"
    ROJO_F  = "C62828"; ROJO_BG  = "FFEBEE"
    thin    = Side(style="thin", color="BDBDBD")
    brd     = Border(left=thin, right=thin, top=thin, bottom=thin)

    def _celda(ws, fila, col, valor, bold=False, bg=None,
               fg="212121", align="left", num_fmt=None):
        c = ws.cell(row=fila, column=col, value=valor)
        c.font      = Font(bold=bold, color=fg, name="Calibri", size=10)
        c.alignment = Alignment(horizontal=align, vertical="center",
                                wrap_text=True)
        c.border    = brd
        if bg:
            c.fill  = PatternFill("solid", fgColor=bg)
        if num_fmt:
            c.number_format = num_fmt
        return c

    # Título
    ws.merge_cells("A1:I2")
    c_t = ws["A1"]
    c_t.value     = "ORTHOCLINIC — Reporte Financiero Total"
    c_t.font      = Font(bold=True, size=14, color="FFFFFF", name="Calibri")
    c_t.fill      = PatternFill("solid", fgColor=AZUL_F)
    c_t.alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[1].height = 24
    ws.row_dimensions[2].height = 24

    ws.merge_cells("A3:I3")
    sub = ws["A3"]
    desc_txt = filtros_desc or "Todos los especialistas / pacientes"
    if saldo_minimo > 0:
        desc_txt += f"   |   Saldo mín.: $ {saldo_minimo:,.2f}"
    sub.value = (f"Generado: {hoy.strftime('%d/%m/%Y %H:%M')}   |   "
                 f"Filtros: {desc_txt}")
    sub.font      = Font(italic=True, size=9, color="616161", name="Calibri")
    sub.alignment = Alignment(horizontal="left", vertical="center")

    COLS = [
        ("Paciente",       22), ("DNI",            13),
        ("Obra Social",    16), ("Teléfono",        14),
        ("Especialista",   22), ("N° Tratamientos",  16),
        ("Total Costo",    14), ("Total Pagado",     14),
        ("Saldo",          14),
    ]
    HDR_ROW = 4
    for i, (lbl, ancho) in enumerate(COLS, 1):
        _celda(ws, HDR_ROW, i, lbl, bold=True,
               bg=AZUL_F, fg="FFFFFF", align="center")
        ws.column_dimensions[get_column_letter(i)].width = ancho
    ws.row_dimensions[HDR_ROW].height = 20

    # Filtrar por saldo mínimo y escribir datos
    datos_filtrados = [d for d in datos if d.get("saldo", 0) >= saldo_minimo]
    gran_costo = gran_pagado = gran_saldo = 0.0

    for i, d in enumerate(datos_filtrados):
        fila = HDR_ROW + 1 + i
        pac  = d.get("paciente") or {}
        esp  = d.get("especialista") or {}
        nom_pac = f"{_fmt(pac.get('apellido',''))} {_fmt(pac.get('nombre',''))}".strip()
        nom_esp = f"{_fmt(esp.get('apellido',''))} {_fmt(esp.get('nombre',''))}".strip()
        costo  = float(d.get("total_costo", 0))
        pagado = float(d.get("total_pagado", 0))
        saldo  = float(d.get("saldo", 0))
        gran_costo  += costo
        gran_pagado += pagado
        gran_saldo  += saldo
        alt = "#FAFAFA" if i % 2 == 1 else "FFFFFF"

        _celda(ws, fila, 1, nom_pac,                     bg=alt)
        _celda(ws, fila, 2, _fmt(pac.get("dni","")),     bg=alt)
        _celda(ws, fila, 3, _fmt(pac.get("obra_social","")), bg=alt)
        _celda(ws, fila, 4, _fmt(pac.get("telefono","")),bg=alt)
        _celda(ws, fila, 5, nom_esp,                     bg=alt)
        _celda(ws, fila, 6, d.get("n_tratamientos", 0),
               bg=alt, align="center")
        _celda(ws, fila, 7, costo,  bg=alt, align="right",
               num_fmt='"$"#,##0.00')
        _celda(ws, fila, 8, pagado, bg=alt, align="right",
               num_fmt='"$"#,##0.00')
        s_bg = ROJO_BG if saldo > 0 else VERDE_BG
        s_fg = ROJO_F  if saldo > 0 else VERDE
        _celda(ws, fila, 9, saldo, bold=True, bg=s_bg, fg=s_fg,
               align="right", num_fmt='"$"#,##0.00')
        ws.row_dimensions[fila].height = 16

    # Totales
    f_tot = HDR_ROW + 1 + len(datos_filtrados)
    ws.merge_cells(f"A{f_tot}:F{f_tot}")
    _celda(ws, f_tot, 1,
           f"TOTALES  ({len(datos_filtrados)} filas)",
           bold=True, bg=AZUL_F, fg="FFFFFF", align="right")
    _celda(ws, f_tot, 7, gran_costo,  bold=True, bg=AZUL_BG,
           align="right", num_fmt='"$"#,##0.00')
    _celda(ws, f_tot, 8, gran_pagado, bold=True, bg=VERDE_BG,
           fg=VERDE, align="right", num_fmt='"$"#,##0.00')
    _celda(ws, f_tot, 9, gran_saldo,  bold=True,
           bg=ROJO_BG if gran_saldo > 0 else VERDE_BG,
           fg=ROJO_F  if gran_saldo > 0 else VERDE,
           align="right", num_fmt='"$"#,##0.00')
    ws.row_dimensions[f_tot].height = 20

    ws.freeze_panes = f"A{HDR_ROW + 1}"
    wb.save(ruta)
    return ruta


# ═══════════════════════════════════════════════════════════════════════════════
#  EXCEL — Agenda Consolidada (multi-especialista, rango libre de fechas)
# ═══════════════════════════════════════════════════════════════════════════════

def generar_excel_agenda_consolidada(datos: list[dict],
                                      especialistas_nombres: list[str] | None = None,
                                      fecha_inicio: str = "",
                                      fecha_fin: str = "",
                                      output_dir: str | None = None) -> str:
    """
    Genera Excel con agenda consolidada de múltiples especialistas.
    datos: resultado de obtener_agenda_consolidada().
    """
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter

    output_dir = output_dir or _OUT_DIR
    os.makedirs(output_dir, exist_ok=True)
    hoy  = datetime.datetime.today()
    ruta = os.path.join(output_dir,
                        f"AgendaConsolidada_{hoy.strftime('%Y%m%d_%H%M')}.xlsx")

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Agenda Consolidada"

    AZUL_F = "1565C0"
    thin   = Side(style="thin", color="BDBDBD")
    brd    = Border(left=thin, right=thin, top=thin, bottom=thin)
    DIAS_ES = {0:"Lunes",1:"Martes",2:"Miércoles",3:"Jueves",
               4:"Viernes",5:"Sábado",6:"Domingo"}
    _EST_BG = {"pendiente": "FFF8E1", "confirmada": "E8F5E9",
               "realizada":  "E3F2FD", "cancelada":  "FFEBEE"}
    _EST_FG = {"pendiente": "E65100", "confirmada": "1B5E20",
               "realizada":  "1565C0", "cancelada":  "C62828"}

    def _celda(ws, fila, col, valor, bold=False, bg=None,
               fg="212121", align="left"):
        c = ws.cell(row=fila, column=col, value=valor)
        c.font      = Font(bold=bold, color=fg, name="Calibri", size=10)
        c.alignment = Alignment(horizontal=align, vertical="center",
                                wrap_text=True)
        c.border    = brd
        if bg:
            c.fill  = PatternFill("solid", fgColor=bg)
        return c

    # Título
    ws.merge_cells("A1:I2")
    c_t = ws["A1"]
    c_t.value     = "ORTHOCLINIC — Agenda Consolidada"
    c_t.font      = Font(bold=True, size=14, color="FFFFFF", name="Calibri")
    c_t.fill      = PatternFill("solid", fgColor=AZUL_F)
    c_t.alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[1].height = 22
    ws.row_dimensions[2].height = 22

    ws.merge_cells("A3:I3")
    c_s = ws["A3"]
    esp_txt = (", ".join(especialistas_nombres)
               if especialistas_nombres else "Todos los especialistas")
    c_s.value = (
        f"Especialistas: {esp_txt}"
        f"   |   Período: {fecha_inicio} → {fecha_fin}"
        f"   |   Generado: {hoy.strftime('%d/%m/%Y %H:%M')}"
    )
    c_s.font  = Font(italic=True, size=9, color="616161", name="Calibri")
    c_s.alignment = Alignment(horizontal="left", vertical="center")

    COLS = [
        ("Fecha",       12), ("Día",          10),
        ("Hora",         8), ("Paciente",     24),
        ("Teléfono",    14), ("Especialista", 22),
        ("Motivo",      28), ("Estado",       14),
        ("Notas",       22),
    ]
    HDR_ROW = 4
    for i, (lbl, ancho) in enumerate(COLS, 1):
        _celda(ws, HDR_ROW, i, lbl, bold=True,
               bg=AZUL_F, fg="FFFFFF", align="center")
        ws.column_dimensions[get_column_letter(i)].width = ancho
    ws.row_dimensions[HDR_ROW].height = 20

    for i, c in enumerate(datos):
        fila   = HDR_ROW + 1 + i
        pac    = c.get("pacientes") or {}
        esp    = c.get("especialistas") or {}
        estado = c.get("estado", "")
        bg     = _EST_BG.get(estado, "FFFFFF")
        fg     = _EST_FG.get(estado, "212121")
        nom_pac = f"{_fmt(pac.get('apellido',''))} {_fmt(pac.get('nombre',''))}".strip()
        nom_esp = f"{_fmt(esp.get('apellido',''))} {_fmt(esp.get('nombre',''))}".strip()
        iso = c.get("fecha_hora", "")
        try:
            dt    = datetime.datetime.fromisoformat(iso.replace("Z", "+00:00"))
            fecha = dt.strftime("%d/%m/%Y")
            dia   = DIAS_ES.get(dt.weekday(), "")
            hora  = dt.strftime("%H:%M")
        except Exception:
            fecha = iso[:10]; dia = ""; hora = ""

        _celda(ws, fila, 1, fecha,   bg=bg, align="center")
        _celda(ws, fila, 2, dia,     bg=bg, align="center")
        _celda(ws, fila, 3, hora,    bg=bg, align="center")
        _celda(ws, fila, 4, nom_pac, bg=bg)
        _celda(ws, fila, 5, _fmt(pac.get("telefono", "")), bg=bg)
        _celda(ws, fila, 6, nom_esp, bg=bg)
        _celda(ws, fila, 7, _fmt(c.get("motivo", "")),     bg=bg)
        _celda(ws, fila, 8, estado.capitalize(), bold=True,
               bg=bg, fg=fg, align="center")
        _celda(ws, fila, 9, _fmt(c.get("notas", "")),      bg=bg)
        ws.row_dimensions[fila].height = 16

    # Total
    f_tot = HDR_ROW + 1 + len(datos)
    ws.merge_cells(f"A{f_tot}:I{f_tot}")
    c_tot = ws.cell(row=f_tot, column=1,
                    value=f"Total: {len(datos)} cita(s)")
    c_tot.font      = Font(bold=True, size=10, color="FFFFFF", name="Calibri")
    c_tot.fill      = PatternFill("solid", fgColor=AZUL_F)
    c_tot.alignment = Alignment(horizontal="right", vertical="center")
    c_tot.border    = brd
    ws.row_dimensions[f_tot].height = 18

    ws.freeze_panes = f"A{HDR_ROW + 1}"
    wb.save(ruta)
    return ruta
